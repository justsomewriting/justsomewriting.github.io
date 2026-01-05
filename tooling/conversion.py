import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
from jinja2 import Template
import os

class DocxToHtmlConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("DOCX to HTML Converter")
        self.root.geometry("700x750")
        
        # Variables
        self.input_path = tk.StringVar()
        self.template_path = tk.StringVar(value="template.html")
        self.post_number = tk.StringVar()
        self.image_width = tk.StringVar(value="554")
        self.image_height = tk.StringVar(value="335")
        self.art_credit = tk.StringVar()
        
        # Music recommendation (optional)
        self.music_title = tk.StringVar()
        self.music_musician = tk.StringVar()

        
        self.create_widgets()
    
    def create_widgets(self):
        # Input File
        tk.Label(self.root, text="Input DOCX File:", font=("Arial", 10, "bold")).grid(
            row=0, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.input_path, width=50).grid(
            row=0, column=1, padx=10, pady=5
        )
        tk.Button(self.root, text="Browse", command=self.browse_input).grid(
            row=0, column=2, padx=10, pady=5
        )
        
        # Template File
        tk.Label(self.root, text="Template HTML File:", font=("Arial", 10, "bold")).grid(
            row=1, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.template_path, width=50).grid(
            row=1, column=1, padx=10, pady=5
        )
        tk.Button(self.root, text="Browse", command=self.browse_template).grid(
            row=1, column=2, padx=10, pady=5
        )
        
        # Post Number
        tk.Label(self.root, text="Post Number:", font=("Arial", 10, "bold")).grid(
            row=2, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.post_number, width=20).grid(
            row=2, column=1, sticky="w", padx=10, pady=5
        )
        tk.Label(self.root, text="(Output: ../posts/[number].html, Image: ../img/img-[number].jpg, Audio: ../audio/[number].mp3)", 
                 font=("Arial", 8), fg="gray").grid(
            row=3, column=0, columnspan=3, sticky="w", padx=10
        )
        
        # Separator
        tk.Frame(self.root, height=2, bg="gray").grid(
            row=4, column=0, columnspan=3, sticky="ew", padx=10, pady=10
        )
        
        # Music Recommendation (Optional)
        tk.Label(self.root, text="Music Recommendation (Optional)", 
                 font=("Arial", 10, "bold")).grid(
            row=5, column=0, columnspan=3, sticky="w", padx=10, pady=(10, 5)
        )
        
        tk.Label(self.root, text="Song Title:", font=("Arial", 10)).grid(
            row=7, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.music_title, width=50).grid(
            row=7, column=1, padx=10, pady=5
        )
        
        tk.Label(self.root, text="Musician:", font=("Arial", 10)).grid(
            row=8, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.music_musician, width=50).grid(
            row=8, column=1, padx=10, pady=5
        )
        
        # Separator
        tk.Frame(self.root, height=2, bg="gray").grid(
            row=9, column=0, columnspan=3, sticky="ew", padx=10, pady=10
        )
        
        # Image settings
        tk.Label(self.root, text="Image Width:", font=("Arial", 10)).grid(
            row=10, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.image_width, width=20).grid(
            row=10, column=1, sticky="w", padx=10, pady=5
        )
        
        tk.Label(self.root, text="Image Height:", font=("Arial", 10)).grid(
            row=11, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.image_height, width=20).grid(
            row=11, column=1, sticky="w", padx=10, pady=5
        )

        tk.Label(self.root, text="Art Credit", font=("Arial", 10)).grid(
            row=12, column=0, sticky="w", padx=10, pady=5
        )
        tk.Entry(self.root, textvariable=self.art_credit, width=20).grid(
            row=12, column=1, sticky="w", padx=10, pady=5
        )
        
        # Convert Button
        tk.Button(
            self.root, text="Convert", command=self.convert, 
            bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), 
            width=20, height=2
        ).grid(row=13, column=0, columnspan=3, pady=20)
        
        # Log/Status Area
        tk.Label(self.root, text="Status Log:", font=("Arial", 10, "bold")).grid(
            row=14, column=0, sticky="w", padx=10, pady=5
        )
        self.log_text = scrolledtext.ScrolledText(
            self.root, height=10, width=80, state="disabled"
        )
        self.log_text.grid(row=15, column=0, columnspan=3, padx=10, pady=5)
    
    def browse_input(self):
        filename = filedialog.askopenfilename(
            title="Select Input DOCX File",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if filename:
            self.input_path.set(filename)
            self.log(f"Input file selected: {filename}")
    
    def browse_template(self):
        filename = filedialog.askopenfilename(
            title="Select Template HTML File",
            filetypes=[("HTML Files", "*.html"), ("All Files", "*.*")]
        )
        if filename:
            self.template_path.set(filename)
            self.log(f"Template file selected: {filename}")
    
    def browse_output(self):
        # No longer needed - removed
        pass
    
    def log(self, message):
        self.log_text.config(state="normal")
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.config(state="disabled")
    
    def convert(self):
        # Validate inputs
        if not self.input_path.get():
            messagebox.showerror("Error", "Please select an input DOCX file")
            return
        if not self.template_path.get():
            messagebox.showerror("Error", "Please select a template HTML file")
            return
        if not self.post_number.get():
            messagebox.showerror("Error", "Please enter a post number")
            return
        
        # Construct output paths
        post_num = self.post_number.get()
        output_path = f"../posts/{post_num}.html"
        image_path = f"../img/img-{post_num}.jpg"
        audio_path = f"../audio/{post_num}.mp3"
        
        try:
            self.log("\n--- Starting Conversion ---")
            self.log(f"Post number: {post_num}")
            self.log(f"Output will be saved to: {output_path}")
            self.log(f"Image path will be: {image_path}")
            
            # Extract from DOCX
            self.log("Reading DOCX file...")
            doc = Document(self.input_path.get())
            
            # Get title (first paragraph)
            title = doc.paragraphs[0].text if doc.paragraphs else "Untitled"
            self.log(f"Title extracted: {title}")
            
            # Get main content
            self.log("Extracting content...")
            content_paragraphs = []
            for para in doc.paragraphs[1:]:
                para_html = ""
                for run in para.runs:
                    text = run.text
                    if run.bold:
                        text = f"<strong>{text}</strong>"
                    if run.italic:
                        text = f"<em>{text}</em>"
                    para_html += text
                if para_html.strip():  # Only add non-empty paragraphs
                    content_paragraphs.append(para_html)
            
            content = "<br><br>".join(content_paragraphs)
            self.log(f"Extracted {len(content_paragraphs)} paragraphs")
            
            # Load template
            self.log("Loading template...")
            with open(self.template_path.get(), 'r', encoding='utf-8') as f:
                template = Template(f.read())
            
            # Render template
            self.log("Rendering HTML...")
            
            # Check if music recommendation is provided
            has_music = bool(self.music_title.get() and self.music_musician.get())
            
            html_output = template.render(
                title=title,
                content=content,
                image_path=image_path,
                image_width=self.image_width.get(),
                image_height=self.image_height.get(),
                has_music=has_music,
                music_path=audio_path if has_music else None,
                music_title=self.music_title.get() if has_music else None,
                music_musician=self.music_musician.get() if has_music else None,
                art_credit=self.art_credit.get()
            )
            
            # Create output directory if it doesn't exist
            os.makedirs("../posts", exist_ok=True)
            
            # Save output
            self.log("Saving output file...")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_output)
            
            self.log(f"✓ Conversion complete! File saved to: {output_path}")
            messagebox.showinfo("Success", f"Conversion completed successfully!\n\nOutput: {output_path}\nImage path: {image_path}")
            
        except Exception as e:
            self.log(f"✗ Error: {str(e)}")
            messagebox.showerror("Error", f"Conversion failed:\n{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = DocxToHtmlConverter(root)
    root.mainloop()